from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX, WD_UNDERLINE

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
# 设置加粗
p.add_run(' bold ').bold = True
p.add_run(' and some ')

# 设置字体
p.add_run(' font name ').font.name = 'Arial'

# 设置字体颜色
p.add_run(' color\n ').font.color.rgb = RGBColor(0xff, 0x99, 0xcc)
p.add_run(' test colors\n ').font.color.rgb = RGBColor(255, 0, 0)
p.add_run(' test colors1\n ').font.color.rgb = RGBColor(0, 255, 0)
p.add_run(' test colors2\n ').font.color.rgb = RGBColor(0, 0, 255)
p.add_run(' test colors3\n ').font.color.rgb = RGBColor(255, 255, 0)
p.add_run(' test colors4\n ').font.color.rgb = RGBColor(0, 255, 255)
p.add_run(' test colors5\n ').font.color.rgb = RGBColor(255, 0, 255)

# 设置字体高亮(背景底色）
# 可设置这些值AUTO = 'default'
# BLACK = 'black'
# BLUE = 'blue'
# BRIGHTGREEN = 'green'
# DARKBLUE = 'darkBlue'
# D