import random

from docx import Document
from docx.shared import RGBColor

document: object = Document()
p = document.add_paragraph('')
yanse = [RGBColor(160, 32, 240), RGBColor(255, 215, 0), RGBColor(0, 0, 0), RGBColor(255, 0, 0), RGBColor(0, 255, 0),
         RGBColor(0, 255, 255)]
text = ["紫 ", "黄 ", "黑 ", "红 ", "绿 ", "蓝 "]

# dayin是随机颜色
dayin: list[int] = []
shuliang: int = 72
j = []
# 把dayin填充够shuliang个随机数
for i in range(shuliang):
    j.append(random.randint(0, 5))
    while j[i] == j[i - 1] and i > 1:
        j[i] = random.randint(0, 5)
    dayin.append(j[i])
# 在文件里输出随机颜色的随机字
for i in dayin:
    p.add_run(random.choice(text)).font.color.rgb = yanse[i]

p = document.add_paragraph('')
# 在文件里输出答案
for i in dayin:
    p.add_run(text[i]).font.color.rgb = yanse[i]

document.save('/home/lmq/Desktop/读颜色游戏.docx')
