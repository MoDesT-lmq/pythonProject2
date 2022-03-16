import random
from docx import Document

weishu = 1
suijishu = []
document: object = Document()
p = document.add_paragraph('')

shuliang = input("请输入数量：")
while not shuliang.isdigit():
    shuliang = input("输入错误，请重新输入数量（只能是数字）：")
shuliang = int(shuliang)
weishu = input("请输入位数：")
while not weishu.isdigit():
    weishu = input("输入错误，请重新输入位数（只能是数字）：")
weishu = int(weishu) #这是多少位
for i in range(shuliang):
    for j in range(weishu):      
        suijishu.append(random.choice(range(10)))
        while suijishu[weishu * i + j] == suijishu[weishu * i + j- 1] and j > 0:
            suijishu[weishu * i + j] = random.choice(range(10))
        p.add_run(str(suijishu[weishu * i + j]))
    p = document.add_paragraph('')

document.save('/home/lmq/Desktop/数字倒读.docx')
